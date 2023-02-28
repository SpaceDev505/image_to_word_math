# import aspose.words as aw
from docx  import Document
from docx.shared import Inches, Cm
from PIL import Image
import pytesseract
import cv2
import os

image = cv2.imread('1.jpg')
gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
cv2.imwrite('gray.jpg', gray)
x = Image.open('gray.jpg')
text = pytesseract.image_to_string(x)
os.remove('gray.jpg')


# doc = aw.Document()

# ocr_doc = aw.DocumentBuilder(doc)
# ocr_doc.write(text)
# doc.save('out.docx')
document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.5)
document.add_paragraph(text)
document.save('out.docx')